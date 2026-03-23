"""Create a minimal Word template and run a post-refactor replacement self-check."""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT.parent) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT.parent))

from independent_case_pipeline.backend.app.config import DEFAULT_REPLACE_MAP_CONFIG
from independent_case_pipeline.backend.app.services.render_service import render_word_from_replace_map
from independent_case_pipeline.backend.app.services.replace_map_service import build_replace_map, read_json, write_replace_map


DEFAULT_SAMPLE_ROOT = PROJECT_ROOT / 'cases'
DEFAULT_DEFANDENT_JSON = DEFAULT_SAMPLE_ROOT / '企业报告.json'
DEFAULT_DEMANDLETTER_JSON = DEFAULT_SAMPLE_ROOT / '律师函.json'
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / 'backend' / 'storage' / 'temp' / 'self_check'


def create_docx_template(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading('Word Replace Self Check', level=1)
    doc.add_paragraph('被告：[--需要替换的被告--]')
    doc.add_paragraph('统一社会信用代码：[--需要替换的统一社会信用代码--]')
    doc.add_paragraph('住所地/送达地址：[--需要替换的住所地/送达地址--]')
    doc.add_paragraph('法定代表人：[--需要替换的法定代表人--]')
    doc.add_paragraph('侵权事实：[--需要替换的侵权事实--]')
    doc.save(path)
    return path


def convert_docx_to_doc(src_docx: Path, dst_doc: Path) -> tuple[bool, str | None]:
    dst_doc.parent.mkdir(parents=True, exist_ok=True)
    errors: list[str] = []

    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except Exception as exc:
        return False, f'pywin32 unavailable: {exc}'

    for progid in ('Word.Application', 'kwps.Application'):
        app = None
        document = None
        pythoncom.CoInitialize()
        try:
            app = win32com.client.Dispatch(progid)
            app.Visible = False
            document = app.Documents.Open(str(src_docx))
            document.SaveAs(str(dst_doc), FileFormat=0)
            return True, None
        except Exception as exc:
            errors.append(f'{progid}: {exc}')
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if app is not None:
                try:
                    app.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    return False, ' | '.join(errors) if errors else 'unknown error'


def read_docx_text(path: Path) -> list[str]:
    doc = Document(str(path))
    return [paragraph.text for paragraph in doc.paragraphs]


def run_self_check(
    *,
    defandent_json: Path,
    demand_letter_json: Path,
    replace_map_config: Path,
    output_root: Path,
) -> dict[str, Any]:
    output_root.mkdir(parents=True, exist_ok=True)

    replace_map = build_replace_map(
        defandent=read_json(defandent_json),
        demand_letter=read_json(demand_letter_json),
        config_path=replace_map_config,
    )
    replace_map_path = write_replace_map(output_root / 'replace_map.json', replace_map)

    template_docx = create_docx_template(output_root / 'self_check_template.docx')
    template_doc = output_root / 'self_check_template.doc'
    doc_created, doc_error = convert_docx_to_doc(template_docx, template_doc)

    docx_render_dir = output_root / 'render_from_docx'
    docx_result = render_word_from_replace_map(
        replace_map=replace_map,
        input_files=[str(template_docx)],
        output_dir=str(docx_render_dir),
        input_base_dir=str(output_root),
        image_align='left',
    )
    rendered_docx = docx_render_dir / template_docx.name
    rendered_docx_text = read_docx_text(rendered_docx) if rendered_docx.is_file() else []

    doc_result: dict[str, Any] | None = None
    rendered_from_doc_text: list[str] = []
    rendered_from_doc_path: Path | None = None
    if doc_created and template_doc.is_file():
        doc_render_dir = output_root / 'render_from_doc'
        doc_result = render_word_from_replace_map(
            replace_map=replace_map,
            input_files=[str(template_doc)],
            output_dir=str(doc_render_dir),
            input_base_dir=str(output_root),
            image_align='left',
        )
        rendered_from_doc_path = doc_render_dir / template_doc.with_suffix('.docx').name
        if rendered_from_doc_path.is_file():
            rendered_from_doc_text = read_docx_text(rendered_from_doc_path)

    report = {
        'status': 'ok',
        'replace_map_json': str(replace_map_path),
        'template_docx': str(template_docx),
        'template_doc': str(template_doc) if template_doc.exists() else None,
        'doc_conversion_success': doc_created,
        'doc_conversion_error': doc_error,
        'render_from_docx': {
            'result': docx_result,
            'output_file': str(rendered_docx),
            'paragraphs': rendered_docx_text,
        },
        'render_from_doc': {
            'result': doc_result,
            'output_file': str(rendered_from_doc_path) if rendered_from_doc_path else None,
            'paragraphs': rendered_from_doc_text,
        },
    }
    report_path = output_root / 'self_check_report.json'
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding='utf-8')
    report['report_path'] = str(report_path)
    return report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Create a temporary Word template and run a replacement self-check.')
    parser.add_argument('--defandent-json', default=str(DEFAULT_DEFANDENT_JSON), help='Path to Defandent.json')
    parser.add_argument('--demand-letter-json', default=str(DEFAULT_DEMANDLETTER_JSON), help='Path to DemandLetter.json')
    parser.add_argument('--replace-map-config', default=str(DEFAULT_REPLACE_MAP_CONFIG), help='Path to replace_map_config.json')
    parser.add_argument('--output-root', default=str(DEFAULT_OUTPUT_ROOT), help='Directory where self-check files will be written')
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    report = run_self_check(
        defandent_json=Path(args.defandent_json).expanduser().resolve(),
        demand_letter_json=Path(args.demand_letter_json).expanduser().resolve(),
        replace_map_config=Path(args.replace_map_config).expanduser().resolve(),
        output_root=Path(args.output_root).expanduser().resolve(),
    )
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == '__main__':
    raise SystemExit(main())

