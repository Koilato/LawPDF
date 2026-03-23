"""Replace placeholders in Word documents using JSON-defined rules."""

from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
import tempfile
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
}


@dataclass
class Replacement:
    keyword: str
    text: str = ""
    images: list[str] = field(default_factory=list)
    image_width_cm: float | None = None
    image_height_cm: float | None = None


@dataclass
class ImageSettings:
    width_cm: float | None = None
    height_cm: float | None = None
    align: str | None = None


@dataclass
class Job:
    input_files: list[str]
    output_dir: str
    input_base_dir: str | None = None
    output_name: str | None = None
    replacements: list[Replacement] = field(default_factory=list)
    image_settings: ImageSettings = field(default_factory=ImageSettings)


def parse_float(value: Any) -> float | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return None


def resolve_alignment(align: str | None):
    if not align:
        return None
    return ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def convert_doc_to_docx(path: str) -> tuple[str, str | None]:
    temp_dir = tempfile.mkdtemp(prefix="word_replace_")
    target = os.path.join(temp_dir, Path(path).stem + ".docx")

    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except Exception as exc:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise RuntimeError("处理 .doc 需要安装 pywin32，且本机安装了 Microsoft Word 或 WPS。") from exc

    errors: list[str] = []
    for progid in ("Word.Application", "kwps.Application"):
        app = None
        doc = None
        pythoncom.CoInitialize()
        try:
            app = win32com.client.Dispatch(progid)
            app.Visible = False
            doc = app.Documents.Open(path)
            doc.SaveAs(target, FileFormat=16)
            return target, temp_dir
        except Exception as exc:
            errors.append(f"{progid}: {exc}")
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if app is not None:
                try:
                    app.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    shutil.rmtree(temp_dir, ignore_errors=True)
    joined = " | ".join(errors) if errors else "未知错误"
    raise RuntimeError(f".doc 转 .docx 失败：{path}；尝试结果：{joined}")


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _copy_run_format(source_run, target_run) -> None:
    if source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))
    if source_run.style is not None:
        target_run.style = source_run.style


def _insert_run_after(run, text: str = "", format_from=None):
    paragraph = run._parent
    new_run = paragraph.add_run()
    run._r.addnext(new_run._r)
    _copy_run_format(format_from or run, new_run)
    if text:
        new_run.text = text
    return new_run


def _find_replacement_matches(text: str, replacements: list[Replacement]) -> list[tuple[int, int, Replacement]]:
    matches: list[tuple[int, int, Replacement]] = []
    index = 0
    while index < len(text):
        next_pos = None
        next_rep = None
        for rep in replacements:
            pos = text.find(rep.keyword, index)
            if pos == -1:
                continue
            if next_pos is None or pos < next_pos or (pos == next_pos and len(rep.keyword) > len(next_rep.keyword)):
                next_pos = pos
                next_rep = rep
        if next_pos is None or next_rep is None:
            break
        matches.append((next_pos, next_pos + len(next_rep.keyword), next_rep))
        index = next_pos + len(next_rep.keyword)
    return matches


def _build_run_ranges(paragraph) -> list[dict[str, Any]]:
    ranges: list[dict[str, Any]] = []
    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        start = cursor
        end = cursor + len(text)
        ranges.append({"index": idx, "run": run, "start": start, "end": end, "text": text})
        cursor = end
    return ranges


def _set_run_text(run, text: str) -> None:
    run.text = text


def _insert_images_after(anchor_run, rep: Replacement, image_settings: ImageSettings, paragraph_align):
    current_anchor = anchor_run
    inserted = False
    for image in rep.images or []:
        if not image:
            continue
        inserted = True
        run = _insert_run_after(current_anchor, format_from=anchor_run)
        run.add_break()
        width_cm = rep.image_width_cm if rep.image_width_cm is not None else image_settings.width_cm
        height_cm = rep.image_height_cm if rep.image_height_cm is not None else image_settings.height_cm
        if width_cm and height_cm:
            run.add_picture(image, width=Cm(width_cm), height=Cm(height_cm))
        elif width_cm:
            run.add_picture(image, width=Cm(width_cm))
        elif height_cm:
            run.add_picture(image, height=Cm(height_cm))
        else:
            run.add_picture(image)
        run.add_break()
        current_anchor = run
    if inserted and paragraph_align is not None:
        anchor_run._parent.alignment = paragraph_align
    return current_anchor


def replace_in_paragraph(paragraph, replacements: Iterable[Replacement], image_settings: ImageSettings, paragraph_align) -> None:
    text = paragraph.text
    if not text:
        return

    replacements = list(replacements)
    matches = _find_replacement_matches(text, replacements)
    if not matches:
        return

    for start, end, rep in reversed(matches):
        run_ranges = _build_run_ranges(paragraph)
        affected = [info for info in run_ranges if info["end"] > start and info["start"] < end]
        if not affected:
            continue

        first = affected[0]
        last = affected[-1]
        first_run = first["run"]
        last_run = last["run"]

        prefix = first["text"][: max(0, start - first["start"])]
        suffix = last["text"][max(0, end - last["start"]):]

        if first["index"] == last["index"]:
            _set_run_text(first_run, prefix)
            anchor = first_run
            if rep.text:
                anchor = _insert_run_after(anchor, rep.text, format_from=first_run)
            anchor = _insert_images_after(anchor, rep, image_settings, paragraph_align)
            if suffix:
                _insert_run_after(anchor, suffix, format_from=first_run)
            continue

        _set_run_text(first_run, prefix)
        for info in affected[1:-1]:
            _set_run_text(info["run"], "")
        _set_run_text(last_run, suffix)

        anchor = first_run
        if rep.text:
            anchor = _insert_run_after(anchor, rep.text, format_from=first_run)
        _insert_images_after(anchor, rep, image_settings, paragraph_align)


def replace_in_document(doc, replacements: Iterable[Replacement], image_settings: ImageSettings, paragraph_align) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements, image_settings, paragraph_align)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements, image_settings, paragraph_align)


def process_documents(job: Job) -> int:
    processed = 0
    input_dir = job.input_base_dir or (os.path.dirname(job.input_files[0]) if job.input_files else "")
    paragraph_align = resolve_alignment(job.image_settings.align)
    if job.output_name and len(job.input_files) != 1:
        raise ValueError("output_name 仅支持单文件渲染")

    for original_src_path in job.input_files:
        if not os.path.isfile(original_src_path):
            continue

        ext = Path(original_src_path).suffix.lower()
        if ext not in {".doc", ".docx"}:
            continue

        working_src_path = original_src_path
        temp_dir: str | None = None
        if ext == ".doc":
            working_src_path, temp_dir = convert_doc_to_docx(original_src_path)

        try:
            rel_path = os.path.relpath(original_src_path, input_dir)
            if rel_path.startswith(".."):
                rel_path = os.path.basename(original_src_path)
        except ValueError:
            rel_path = os.path.basename(original_src_path)

        if job.output_name:
            rel_docx = str(Path(job.output_name).with_suffix(".docx"))
        else:
            rel_docx = str(Path(rel_path).with_suffix(".docx"))
        dst_path = os.path.join(job.output_dir, rel_docx)
        os.makedirs(os.path.dirname(dst_path), exist_ok=True)

        try:
            doc = Document(working_src_path)
            replace_in_document(doc, job.replacements, job.image_settings, paragraph_align)
            doc.save(dst_path)
            processed += 1
        finally:
            if temp_dir:
                shutil.rmtree(temp_dir, ignore_errors=True)

    return processed


def load_job(path: Path) -> Job:
    data = json.loads(path.read_text(encoding="utf-8-sig"))
    input_files = list(data.get("input_files") or [])
    output_dir = str(data.get("output_dir") or "").strip()
    input_base_dir = data.get("input_base_dir")

    replacements: list[Replacement] = []
    for item in data.get("replacements") or []:
        keyword = str(item.get("keyword") or "").strip()
        if not keyword:
            continue
        images = item.get("images")
        if images is None:
            images = item.get("image", [])
        if isinstance(images, str):
            images = [images] if images else []
        replacements.append(
            Replacement(
                keyword=keyword,
                text=str(item.get("text") or ""),
                images=list(images or []),
                image_width_cm=parse_float(item.get("image_width_cm")),
                image_height_cm=parse_float(item.get("image_height_cm")),
            )
        )

    if not input_files:
        raise ValueError("input_files 不能为空")
    if not output_dir:
        raise ValueError("output_dir 不能为空")
    if not replacements:
        raise ValueError("replacements 不能为空")

    image_settings = ImageSettings(
        width_cm=parse_float(data.get("image_width_cm")),
        height_cm=parse_float(data.get("image_height_cm")),
        align=data.get("image_align") or None,
    )

    return Job(
        input_files=input_files,
        output_dir=output_dir,
        input_base_dir=input_base_dir,
        output_name=str(data.get("output_name") or "").strip() or None,
        replacements=replacements,
        image_settings=image_settings,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Replace keywords in doc and docx files using a JSON job file.")
    parser.add_argument("--config", required=True, help="Path to the job JSON file.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    config_path = Path(args.config).expanduser().resolve()
    if not config_path.is_file():
        print(f"MISSING {config_path}", file=sys.stderr)
        return 1

    job = load_job(config_path)
    processed = process_documents(job)
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    print(json.dumps({"status": "ok", "processed": processed, "output_dir": job.output_dir}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

